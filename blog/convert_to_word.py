from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_markdown_to_word(markdown_file, output_file):
    """Convert markdown file to Word document with basic formatting"""
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r')
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Handle bold text
            text = text.replace('**', '')
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            text = text.replace('**', '')
            p = doc.add_paragraph(text, style='List Number')
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip('\r'))
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
        
        # Tables (basic markdown table support)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\r'))
                i += 1
            i -= 1
            
            if len(table_lines) > 2:
                # Parse table
                rows = [row.split('|')[1:-1] for row in table_lines if not row.strip().startswith('|:')]
                rows = [[cell.strip() for cell in row] for row in rows if row]
                
                if rows:
                    # Skip separator row
                    header = rows[0]
                    data_rows = [r for r in rows[2:] if r] if len(rows) > 2 else []
                    
                    table = doc.add_table(rows=len(data_rows) + 1, cols=len(header))
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    for j, cell_text in enumerate(header):
                        table.rows[0].cells[j].text = cell_text
                    
                    # Data rows
                    for row_idx, row_data in enumerate(data_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx + 1].cells):
                                table.rows[row_idx + 1].cells[col_idx].text = cell_text
        
        # Regular paragraphs
        else:
            # Clean up markdown formatting
            text = line
            text = text.replace('**', '')
            text = text.replace('*', '')
            text = text.replace('`', '')
            
            # Remove HTML anchor tags
            text = re.sub(r'<a id="[^"]*"></a>', '', text)
            
            if text.strip():
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Save the document
    doc.save(output_file)
    print(f"✅ Successfully converted to Word document: {output_file}")

# Execute conversion
if __name__ == "__main__":
    markdown_file = r"d:\Manav\premium extension\flowboard_5\blog\blog_agency_lifecycle.md"
    output_file = r"d:\Manav\premium extension\flowboard_5\blog\blog_agency_lifecycle.docx"
    
    convert_markdown_to_word(markdown_file, output_file)
